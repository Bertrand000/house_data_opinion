
from docx import Document
from docx.shared import Inches

def gen_docfile(context="",zongjie="",pie_file_path="",doc_file_path="C:\\test.doc",title="自动分析报告"):
    '''
    保存word分析报告
    :param context: 对于XXXXX部分分析得出以下结论XXX
    :param zongjie: 总结语句
    :param pie_file_path: 需要保存的WORK文件路径
    :return: 无返回值
    '''
    # 新建一个文档
    document = Document()
    document.add_heading(title, 0)
    # 添加一个段落
    document.add_paragraph(context)
    document.add_heading(u'图文分析结论', level=1)

    # 插入图片，文件名可以作为参数传入，由之前的程序进行传入
    document.add_picture(pie_file_path, width=Inches(5.0))
    # document.add_page_break()

    document.add_paragraph(zongjie)

    document.save(doc_file_path)